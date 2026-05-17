import json
import re
import io
import os
from datetime import date, timedelta
from urllib.parse import quote

import requests as http_requests
from PIL import Image as PILImage

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail
from django.db.models import Q, Count
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.http import require_POST, require_GET
from django.views.decorators.csrf import csrf_protect

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.models import CustomUser, Notification, ActivityLog
from core.utils import (
    log_activity, get_client_ip,
    send_approval_email, send_rejection_email,
    create_notification, fmt_dt
)

from beneficiary.models import (
    OrphanForm, OrphanMother, OrphanFather,
    SpecialNeedsForm, FamilyForm, FamilyWife,
    CurrentGuardian, FamilyMember,
    OrphanDocument, SpecialDocument, FamilyDocument,
)

from sponsor.models import SponsorProfile
from .decorators import admin_required

# Optional dependency: openpyxl
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    OPENPYXL = True
except ImportError:
    OPENPYXL = False

from .decorators import admin_required
from django.views.decorators.http import require_GET
from core.models import CustomUser


# ══════════════════════════════════════════════
#  الإحصائيات
# ══════════════════════════════════════════════
def _build_stats():
    from datetime import timedelta

    today    = date.today()
    week_ago = today - timedelta(days=7)

    base_all = CustomUser.objects.filter(
        is_active=True
    ).exclude(user_type='admin')

    def _stat(user_type):
        base = base_all.filter(user_type=user_type)
        return {
            'total':          base.count(),
            'pending':        base.filter(is_approved=False).count(),
            'approved_today': base.filter(
                is_approved=True, date_joined__date=today
            ).count(),
            'rejected_today': ActivityLog.objects.filter(
                action='REJECT',
                created_at__date=today,
                description__icontains=user_type,
            ).count(),
        }

    return {
        # إجمالي
        'total':          base_all.count(),
        'pending':        base_all.filter(is_approved=False).count(),
        'approved_today': base_all.filter(
            is_approved=True, date_joined__date=today
        ).count(),
        'rejected_today': ActivityLog.objects.filter(
            action='REJECT', created_at__date=today
        ).count(),

        # هذا الأسبوع
        'this_week': base_all.filter(
            date_joined__date__gte=week_ago
        ).count(),

        # حسب النوع
        'sponsor': _stat('sponsor'),
        'orphan':  _stat('orphan'),
        'family':  _stat('family'),
        'special': _stat('special'),
    }

# ══════════════════════════════════════════════
#  تجميع بيانات المستخدم التفصيلية
# ══════════════════════════════════════════════
def _user_detail(user):
    """جمع كل البيانات التفصيلية لمستخدم واحد"""
    d  = {}
    ut = user.user_type

    # ── كافل ──
    if ut == 'sponsor':
        try:
            p = SponsorProfile.objects.get(user=user)
            d['sponsor'] = {
                'job':       p.job,
                'country':   p.country,
                'city':      p.city,
                'photo_url': p.photo.url if p.photo else None,
            }
        except SponsorProfile.DoesNotExist:
            pass

    # ── يتيم ──
    elif ut == 'orphan':
        try:
            o = OrphanForm.objects.get(user=user)
            d['orphan'] = {
                'birth_date':        str(o.birth_date),
                'orphan_type':       o.orphan_type,
                'health_status':     o.health_status,
                'education_level':   o.education_level,
                'school_grade':      o.school_grade   or '',
                'school_name':       o.school_name    or '',
                'current_city':      o.current_city,
                'current_street':    o.current_street,
                'current_landmark':  o.current_landmark,
                'previous_city':     o.previous_city  or '',
                'previous_street':   o.previous_street or '',
                'previous_landmark': o.previous_landmark or '',
                'housing_type':      o.housing_type,
                'housing_ownership': o.housing_ownership,
                'monthly_rent':      str(o.monthly_rent) if o.monthly_rent else '',
                'story':             o.story,
                'photo_url':         o.photo.url if o.photo else None,
                # ── التواصل ──
                'phone2':            o.phone2         or '',
                'phone2_country':    o.phone2_country or '+970',
                'whatsapp':          o.whatsapp       or '',
                'whatsapp_country':  o.whatsapp_country or '+970',
            }
        except OrphanForm.DoesNotExist:
            pass

        try:
            m = OrphanMother.objects.get(form__user=user)
            d['mother'] = {
                'full_name':       m.get_full_name(),
                'id_number':       m.id_number,
                'birth_date':      str(m.birth_date),
                'is_alive':        m.is_alive,
                'death_date':      str(m.death_date) if m.death_date else '',
                'death_reason':    m.death_reason    or '',
                'health_status':   m.health_status,
                'education_level': m.education_level,
                'job':             m.job,
                'monthly_income':  str(m.monthly_income),
            }
        except OrphanMother.DoesNotExist:
            pass

        try:
            f = OrphanFather.objects.get(form__user=user)
            d['father'] = {
                'full_name':       f.get_full_name(),
                'id_number':       f.id_number,
                'birth_date':      str(f.birth_date),
                'is_alive':        f.is_alive,
                'death_date':      str(f.death_date) if f.death_date else '',
                'death_reason':    f.death_reason    or '',
                'health_status':   f.health_status,
                'education_level': f.education_level,
                'job':             f.job,
                'children_count':  f.children_count,
                'income_before':   str(f.income_before) if f.income_before else '',
                'pension_after':   str(f.pension_after)  if f.pension_after  else '',
            }
        except OrphanFather.DoesNotExist:
            pass

        docs      = OrphanDocument.objects.filter(form__user=user)
        d['docs'] = [{'type': x.doc_type, 'url': x.file.url, 'name': x.file.name} for x in docs]

    # ── ذوو احتياجات ──
    elif ut == 'special':
        try:
            s = SpecialNeedsForm.objects.get(user=user)
            d['special'] = {
                'birth_date':        str(s.birth_date),
                'health_status':     s.health_status,
                'education_level':   s.education_level,
                'school_grade':      s.school_grade   or '',
                'school_name':       s.school_name    or '',
                'current_city':      s.current_city,
                'current_street':    s.current_street,
                'current_landmark':  s.current_landmark,
                'previous_city':     s.previous_city  or '',
                'previous_street':   s.previous_street or '',
                'previous_landmark': s.previous_landmark or '',
                'housing_type':      s.housing_type,
                'housing_ownership': s.housing_ownership,
                'monthly_rent':      str(s.monthly_rent) if s.monthly_rent else '',
                'case_details':      s.case_details,
                'photo_url':         s.photo.url if s.photo else None,
                # ── التواصل ──
                'phone2':            s.phone2         or '',
                'phone2_country':    s.phone2_country or '+970',
                'whatsapp':          s.whatsapp       or '',
                'whatsapp_country':  s.whatsapp_country or '+970',
            }
        except SpecialNeedsForm.DoesNotExist:
            pass

        docs      = SpecialDocument.objects.filter(form__user=user)
        d['docs'] = [{'type': x.doc_type, 'url': x.file.url, 'name': x.file.name} for x in docs]

    # ── أسرة ──
    elif ut == 'family':
        try:
            f = FamilyForm.objects.get(user=user)
            d['family'] = {
                'birth_date':        str(f.birth_date),
                'id_number':         f.id_number,
                'marital_status':    f.marital_status,
                'health_status':     f.health_status,
                'education_level':   f.education_level,
                'job':               f.job,
                'current_city':      f.current_city,
                'current_street':    f.current_street,
                'current_landmark':  f.current_landmark,
                'previous_city':     f.previous_city  or '',
                'previous_street':   f.previous_street or '',
                'previous_landmark': f.previous_landmark or '',
                'housing_type':      f.housing_type,
                'housing_ownership': f.housing_ownership,
                'monthly_rent':      str(f.monthly_rent) if f.monthly_rent else '',
                'members_count':     f.family_members_count,
                'sick_count':        f.sick_members_count,
                'general_status':    f.general_status,
                'photo_url':         f.photo.url if f.photo else None,
                # ── التواصل ──
                'phone2':            f.phone2         or '',
                'phone2_country':    f.phone2_country or '+970',
                'whatsapp':          f.whatsapp       or '',
                'whatsapp_country':  f.whatsapp_country or '+970',
            }
        except FamilyForm.DoesNotExist:
            pass

        try:
            w = FamilyWife.objects.get(form__user=user)
            d['wife'] = {
                'full_name':       w.get_full_name(),
                'id_number':       w.id_number,
                'birth_date':      str(w.birth_date),
                'health_status':   w.health_status,
                'education_level': w.education_level,
                'photo_url':       w.photo.url if w.photo else None,
            }
        except FamilyWife.DoesNotExist:
            pass

        docs      = FamilyDocument.objects.filter(form__user=user)
        d['docs'] = [{'type': x.doc_type, 'url': x.file.url, 'name': x.file.name} for x in docs]

    # ── المعيل ──
    g = CurrentGuardian.objects.filter(
        form_type=ut, form_id=_get_form_id(user)
    ).first()
    if g:
        d['guardian'] = {
            'full_name':       g.get_full_name(),
            'id_number':       g.id_number,
            'gender':          g.gender,
            'relation':        g.relation,
            'job':             g.job,
            'health_status':   g.health_status,
            'education_level': g.education_level,
            'monthly_income':  str(g.monthly_income),
            'dependents':      g.dependents,
        }

    # ── أفراد الأسرة ──
    members      = FamilyMember.objects.filter(form_type=ut, form_id=_get_form_id(user))
    d['members'] = [{
        'full_name':       m.get_full_name(),
        'id_number':       m.id_number,
        'gender':          m.gender,
        'birth_date':      str(m.birth_date),
        'marital_status':  m.marital_status,
        'relation':        m.relation,
        'health_status':   m.health_status,
        'education_level': m.education_level,
    } for m in members]

    # ── سجل النشاط ──
    logs = ActivityLog.objects.filter(
        Q(user=user) | Q(target_id=user.pk, target_model='CustomUser')
    ).order_by('-created_at')[:10]
    d['logs'] = [{
        'description': l.description,
        'created_at':  fmt_dt(l.created_at),
        'user':        l.user.get_short_name() if l.user else 'النظام',
        'icon':        '✅' if 'APPROVE' in l.action else '❌' if 'REJECT' in l.action else '📌',
    } for l in logs]

    # ── الكافل ──
    sponsor_info = None
    if user.is_approved and ut in ['orphan', 'family', 'special']:
        try:
            from beneficiary.models import OrphanForm as OF, SpecialNeedsForm as SNF, FamilyForm as FF
            model_map = {'orphan': OF, 'family': FF, 'special': SNF}
            form      = model_map[ut].objects.get(user=user)
            if hasattr(form, 'sponsor') and form.sponsor:
                sp = form.sponsor.user
                sponsor_info = {
                    'name':       sp.get_full_name(),
                    'reg_number': sp.registration_number or '—',
                    'email':      sp.email,
                    'phone':      f'{sp.phone_country}{sp.phone}',
                }
        except Exception:
            pass
    d['sponsor_info'] = sponsor_info

    # ── المدفوعات ──
    from core.models import Payment, Aid
    STATUS_MAP  = {'paid': 'مدفوعة', 'pending': 'معلّقة', 'late': 'متأخرة'}
    PAID_BY_MAP = {'sponsor': 'كافل', 'admin': 'إدارة', 'external': 'جهة خارجية'}

    payments = []
    for p in Payment.objects.filter(
        beneficiary=user
    ).select_related('sponsor', 'created_by').order_by('-date')[:20]:
        if p.paid_by == 'sponsor' and p.sponsor:
            party = p.sponsor.get_full_name()
        elif p.paid_by == 'admin' and p.created_by:
            party = p.created_by.get_full_name() or p.created_by.username
        elif p.paid_by == 'external':
            party = p.paid_by_note or '—'
        else:
            party = '—'
        payments.append({
            'date':          str(p.date),
            'amount_ils':    str(p.amount_ils),
            'amount_usd':    str(p.amount_usd),
            'paid_by_label': PAID_BY_MAP.get(p.paid_by, p.paid_by),
            'party':         party,
            'status':        p.status,
            'status_label':  STATUS_MAP.get(p.status, p.status),
            'note':          p.note or '',
        })
    d['payments'] = payments

    # ── المساعدات ──
    AID_TYPES = {
        'food': 'غذائية', 'medical': 'طبية', 'financial': 'مالية',
        'clothing': 'ملابس', 'furniture': 'أثاث', 'education': 'تعليمية', 'other': 'أخرى',
    }
    aids = []
    for a in Aid.objects.filter(
        beneficiary=user
    ).select_related('created_by').order_by('-date')[:20]:
        aids.append({
            'date':           str(a.date),
            'name':           a.name,
            'aid_type_label': AID_TYPES.get(a.aid_type, a.aid_type),
            'quantity':       a.quantity,
            'provider':       a.provider,
            'created_by':     (a.created_by.get_full_name() or a.created_by.username) if a.created_by else 'النظام',
            'note':           a.note or '',
        })
    d['aids'] = aids

    # ── بيانات التواصل الإجمالية ──
    form_obj              = d.get('orphan') or d.get('family') or d.get('special') or {}
    d['phone2']           = (form_obj.get('phone2_country', '') or '') + (form_obj.get('phone2', '') or '')
    d['phone2_country']   = form_obj.get('phone2_country', '') or '+970'
    d['whatsapp']         = (user.whatsapp_country or '') + (user.whatsapp or '') if user.whatsapp else ''
    d['whatsapp_country'] = user.whatsapp_country or '+970'
    d['nationality_code'] = user.nationality_code or ''

    return d

def _get_form_id(user):
    """الحصول على معرف النموذج لمستخدم"""
    try:
        if user.user_type == 'orphan':  return OrphanForm.objects.get(user=user).id
        if user.user_type == 'special': return SpecialNeedsForm.objects.get(user=user).id
        if user.user_type == 'family':  return FamilyForm.objects.get(user=user).id
    except Exception:
        pass
    return None


# ══════════════════════════════════════════════
#  الصفحة الرئيسية
# ══════════════════════════════════════════════
@admin_required
def requests_list(request):
    stats = _build_stats()
    notif_count = Notification.objects.filter(
        recipient=request.user, is_read=False
    ).count()
    return render(request, 'admin_panel/requests.html', {
        'stats': stats, 'notif_count': notif_count,
    })


# ══════════════════════════════════════════════
#  API — بيانات الجدول
# ══════════════════════════════════════════════
@admin_required
@require_GET
def requests_data(request):
    """إرجاع كل بيانات المستخدمين المعلّقين مع التفاصيل"""
    users = CustomUser.objects.filter(
        is_approved=False, is_active=True
    ).exclude(user_type='admin').order_by('-date_joined')

    data = []
    for u in users:
        photo = None
        try:
            if u.user_type == 'sponsor':
                p = SponsorProfile.objects.get(user=u)
                photo = p.photo.url if p.photo else None
            elif u.user_type == 'orphan':
                p = OrphanForm.objects.get(user=u)
                photo = p.photo.url if p.photo else None
            elif u.user_type == 'special':
                p = SpecialNeedsForm.objects.get(user=u)
                photo = p.photo.url if p.photo else None
            elif u.user_type == 'family':
                p = FamilyForm.objects.get(user=u)
                photo = p.photo.url if p.photo else None
        except Exception:
            pass

        data.append({
            'id':           str(u.pk),
            'full_name':    u.get_full_name(),
            'user_type':    u.user_type,
            'email':        u.email,
            'phone':        u.phone or '',
            'phone_country': u.phone_country or '',
            'whatsapp':     u.whatsapp or '',
            'whatsapp_country': u.whatsapp_country or '',
            'id_number':    u.id_number or '',
            'nationality':  u.nationality or '',
            'gender':       u.gender or '',
            'date_joined':  u.date_joined.isoformat(),
            'reg_number': u.registration_number or '',
            'photo':        photo,
            'detail':       _user_detail(u),
        })

    return JsonResponse({'users': data, 'total': len(data)})


# ══════════════════════════════════════════════
#  الموافقة
# ══════════════════════════════════════════════
@admin_required
@require_POST
@csrf_protect
def approve_user(request):
    user_id = request.POST.get('user_id', '').strip()
    if not user_id:
        return JsonResponse({'status': 'error', 'message': 'معرف غير صالح'})
    try:
        user = CustomUser.objects.get(pk=user_id, is_approved=False)
        from core.utils import generate_reg_number

        user.is_approved = True
        user.registration_number = generate_reg_number(user)
        user.save(update_fields=['is_approved', 'registration_number'])

        create_notification(
            recipient=user, ntype='APPROVED',
            title='تم قبول حسابك ✅',
            message='تمت الموافقة على حسابك، يمكنك الآن تسجيل الدخول.',
            sender=request.user, action_url='/login/',
        )
        send_approval_email(user)
        log_activity(
            request.user, 'APPROVE',
            description=f'موافقة على حساب: {user.get_full_name()}',
            target_model='CustomUser', target_id=user.pk, request=request,
        )
        return JsonResponse({'status': 'success', 'message': f'تمت الموافقة على {user.get_full_name()} ✅'})

    except CustomUser.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'المستخدم غير موجود أو تمت معالجته'})


# ══════════════════════════════════════════════
#  الرفض
# ══════════════════════════════════════════════
@admin_required
@require_POST
@csrf_protect
def reject_user(request):
    user_id = request.POST.get('user_id', '').strip()
    reason  = request.POST.get('reason', '').strip()
    if not user_id:
        return JsonResponse({'status': 'error', 'message': 'معرف غير صالح'})
    if not reason:
        return JsonResponse({'status': 'error', 'message': 'سبب الرفض مطلوب'})
    try:
        user = CustomUser.objects.get(pk=user_id, is_approved=False)
        name = user.get_full_name()

        send_rejection_email(user, complaint_url='/complaints/')
        log_activity(
            request.user, 'REJECT',
            description=f'رفض حساب: {name} — السبب: {reason}',
            target_model='CustomUser', target_id=user.pk, request=request,
        )
        user.delete()
        return JsonResponse({'status': 'success', 'message': f'تم رفض طلب {name} ❌'})

    except CustomUser.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'المستخدم غير موجود أو تمت معالجته'})


# ══════════════════════════════════════════════
#  إرسال رسالة بريد إلكتروني
# ══════════════════════════════════════════════
@admin_required
@require_POST
@csrf_protect
def send_message(request):
    user_id = request.POST.get('user_id', '').strip()
    subject = request.POST.get('subject', 'رسالة من الإدارة').strip()
    message = request.POST.get('message', '').strip()

    if not user_id or not message:
        return JsonResponse({'status': 'error', 'message': 'بيانات غير مكتملة'})

    try:
        user = CustomUser.objects.get(pk=user_id)

        # إرسال البريد
        send_mail(
            subject=subject,
            message=message,
            from_email=settings.DEFAULT_FROM_EMAIL,
            recipient_list=[user.email],
            fail_silently=False,
        )

        # تخزين في قاعدة البيانات كإشعار
        create_notification(
            recipient=user,
            ntype='MESSAGE',
            title=subject,
            message=message,
            sender=request.user,
        )

        # تسجيل في سجل النشاط
        log_activity(
            request.user, 'MESSAGE',
            description=f'رسالة بريد إلى: {user.get_full_name()} — {subject}',
            target_model='CustomUser',
            target_id=user.pk,
            request=request,
        )

        return JsonResponse({'status': 'success', 'message': 'تم إرسال الرسالة ✅'})

    except CustomUser.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'المستخدم غير موجود'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'فشل الإرسال: {str(e)}'})


# ══════════════════════════════════════════════
#  ملاحظة داخلية
# ══════════════════════════════════════════════
@admin_required
@require_POST
@csrf_protect
def add_note(request):
    user_id = request.POST.get('user_id', '').strip()
    note    = request.POST.get('note', '').strip()

    if not user_id.isdigit() or not note:
        return JsonResponse({'status': 'error', 'message': 'بيانات غير مكتملة'})

    try:
        user = CustomUser.objects.get(pk=user_id)
        log_activity(
            request.user, 'NOTE',
            description=f'ملاحظة داخلية على: {user.get_full_name()} — {note}',
            target_model='CustomUser', target_id=user.pk, request=request,
        )
        return JsonResponse({'status': 'success', 'message': 'تم حفظ الملاحظة 📝'})

    except CustomUser.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'المستخدم غير موجود'})


# ══════════════════════════════════════════════
#  تصدير Excel
# ══════════════════════════════════════════════
@admin_required
@require_GET
def export_requests(request):
    if not OPENPYXL:
        return HttpResponse('مكتبة openpyxl غير مثبتة', status=500)

    ids  = request.GET.get('ids', '')
    mode = request.GET.get('type', 'all')

    # ── جلب المستخدمين ──
    if ids:
        id_list = [i.strip() for i in ids.split(',') if i.strip()]
        users = list(CustomUser.objects.filter(pk__in=id_list))
    else:
        users = list(CustomUser.objects.filter(
            is_approved=False, is_active=True
        ).exclude(user_type='admin').order_by('-date_joined'))

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ── تنسيقات ──
    HDR_FILL  = PatternFill('solid', fgColor='7C3AED')
    HDR_FONT  = Font(color='FFFFFF', bold=True, size=11)
    HDR_ALIGN = Alignment(horizontal='center', vertical='center', wrap_text=True)
    DAT_ALIGN = Alignment(horizontal='right',  vertical='center', wrap_text=True)
    thin      = Side(border_style='thin', color='E2E8F0')
    BORDER    = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _hdr(ws, cols):
        ws.append(cols)
        for cell in ws[1]:
            cell.fill      = HDR_FILL
            cell.font      = HDR_FONT
            cell.alignment = HDR_ALIGN
            cell.border    = BORDER
        ws.row_dimensions[1].height = 28

    def _row(ws, idx):
        fill = PatternFill('solid', fgColor='F5F3FF' if idx % 2 == 0 else 'FFFFFF')
        for cell in ws[idx]:
            cell.fill      = fill
            cell.alignment = DAT_ALIGN
            cell.border    = BORDER

    def _auto_width(ws):
        for col in ws.columns:
            max_w = max((len(str(c.value or '')) for c in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_w + 4, 50)

    TYPE_MAP = {'sponsor': 'كافل', 'orphan': 'يتيم',
                'family': 'أسرة', 'special': 'ذوو احتياجات'}

    # ══ ورقة 1 — البيانات الأساسية ══
    ws1 = wb.create_sheet('البيانات الأساسية')
    _hdr(ws1, ['الاسم الكامل', 'النوع', 'البريد الإلكتروني',
               'الجوال الأول', 'الجوال الثاني', 'الواتساب',
               'رقم الهوية', 'الجنسية', 'الجنس', 'تاريخ التسجيل'])
    for i, u in enumerate(users, 2):
        ws1.append([
            u.get_full_name(),
            TYPE_MAP.get(u.user_type, u.user_type),
            u.email or '',
            f'{u.phone_country or ""}{u.phone or ""}',
            f'{u.phone2_country or ""}{u.phone2 or ""}' if getattr(u, 'phone2', None) else '',
            f'{u.whatsapp_country or ""}{u.whatsapp or ""}' if getattr(u, 'whatsapp', None) else '',
            u.id_number or '',
            u.nationality or '',
            u.gender or '',
            fmt_dt(u.date_joined),
        ])
        _row(ws1, i)
    _auto_width(ws1)

    # ══ ورقة 2 — الأيتام ══
    orphan_users = [u for u in users if u.user_type == 'orphan']
    if orphan_users:
        ws2 = wb.create_sheet('الأيتام')
        _hdr(ws2, ['الاسم الكامل', 'تاريخ الميلاد', 'نوع اليتم',
                   'الحالة الصحية', 'المستوى التعليمي', 'اسم المدرسة',
                   'المدينة الحالية', 'الشارع الحالي', 'أقرب معلم',
                   'نوع السكن', 'ملكية السكن', 'مبلغ الإيجار', 'القصة'])
        for i, u in enumerate(orphan_users, 2):
            try:
                o = OrphanForm.objects.get(user=u)
                ws2.append([
                    u.get_full_name(), str(o.birth_date), o.orphan_type,
                    o.health_status, o.education_level, o.school_name or '',
                    o.current_city, o.current_street or '', o.current_landmark or '',
                    o.housing_type, o.housing_ownership,
                    str(o.monthly_rent) if o.monthly_rent else '',
                    o.story or '',
                ])
            except OrphanForm.DoesNotExist:
                ws2.append([u.get_full_name()] + ['—'] * 12)
            _row(ws2, i)
        _auto_width(ws2)

    # ══ ورقة 3 — الأمهات ══
    if orphan_users:
        ws3 = wb.create_sheet('أمهات الأيتام')
        _hdr(ws3, ['اسم اليتيم', 'اسم الأم', 'رقم الهوية', 'تاريخ الميلاد',
                   'الحالة', 'سبب الوفاة', 'الصحة', 'التعليم', 'المهنة', 'الدخل الشهري'])
        i = 2
        for u in orphan_users:
            try:
                m = OrphanMother.objects.get(form__user=u)
                ws3.append([
                    u.get_full_name(), m.get_full_name(), m.id_number or '',
                    str(m.birth_date),
                    'على قيد الحياة' if m.is_alive else 'متوفية',
                    m.death_reason or '',
                    m.health_status, m.education_level, m.job,
                    str(m.monthly_income),
                ])
            except OrphanMother.DoesNotExist:
                ws3.append([u.get_full_name()] + ['—'] * 9)
            _row(ws3, i); i += 1
        _auto_width(ws3)

    # ══ ورقة 4 — الآباء ══
    if orphan_users:
        ws4 = wb.create_sheet('آباء الأيتام')
        _hdr(ws4, ['اسم اليتيم', 'اسم الأب', 'رقم الهوية', 'تاريخ الميلاد',
                   'الحالة', 'سبب الوفاة', 'الصحة', 'التعليم', 'المهنة', 'عدد الأبناء'])
        i = 2
        for u in orphan_users:
            try:
                f = OrphanFather.objects.get(form__user=u)
                ws4.append([
                    u.get_full_name(), f.get_full_name(), f.id_number or '',
                    str(f.birth_date),
                    'على قيد الحياة' if f.is_alive else 'متوفي',
                    f.death_reason or '',
                    f.health_status, f.education_level, f.job, f.children_count,
                ])
            except OrphanFather.DoesNotExist:
                ws4.append([u.get_full_name()] + ['—'] * 9)
            _row(ws4, i); i += 1
        _auto_width(ws4)

    # ══ ورقة 5 — الأسر ══
    family_users = [u for u in users if u.user_type == 'family']
    if family_users:
        ws5 = wb.create_sheet('الأسر')
        _hdr(ws5, ['الاسم الكامل', 'رقم الهوية', 'تاريخ الميلاد',
                   'الحالة الاجتماعية', 'الصحة', 'التعليم', 'المهنة',
                   'المدينة الحالية', 'الشارع', 'نوع السكن', 'ملكية السكن',
                   'عدد الأفراد', 'عدد المرضى', 'الوضع العام'])
        for i, u in enumerate(family_users, 2):
            try:
                f = FamilyForm.objects.get(user=u)
                ws5.append([
                    u.get_full_name(), f.id_number or '', str(f.birth_date),
                    f.marital_status, f.health_status, f.education_level, f.job,
                    f.current_city, f.current_street or '',
                    f.housing_type, f.housing_ownership,
                    f.family_members_count, f.sick_members_count,
                    f.general_status or '',
                ])
            except FamilyForm.DoesNotExist:
                ws5.append([u.get_full_name()] + ['—'] * 13)
            _row(ws5, i)
        _auto_width(ws5)

    # ══ ورقة 6 — الزوجات ══
    if family_users:
        ws6 = wb.create_sheet('زوجات الأسر')
        _hdr(ws6, ['اسم رب الأسرة', 'اسم الزوجة', 'رقم الهوية',
                   'تاريخ الميلاد', 'الصحة', 'التعليم'])
        i = 2
        for u in family_users:
            try:
                w = FamilyWife.objects.get(form__user=u)
                ws6.append([
                    u.get_full_name(), w.get_full_name(), w.id_number or '',
                    str(w.birth_date), w.health_status, w.education_level,
                ])
            except FamilyWife.DoesNotExist:
                pass
            else:
                _row(ws6, i); i += 1
        if i == 2:
            ws6.append(['لا توجد بيانات زوجات'])
        _auto_width(ws6)

    # ══ ورقة 7 — ذوو الاحتياجات ══
    special_users = [u for u in users if u.user_type == 'special']
    if special_users:
        ws7 = wb.create_sheet('ذوو الاحتياجات')
        _hdr(ws7, ['الاسم الكامل', 'تاريخ الميلاد', 'الحالة الصحية',
                   'المستوى التعليمي', 'المدينة الحالية', 'الشارع',
                   'نوع السكن', 'ملكية السكن', 'تفاصيل الحالة'])
        for i, u in enumerate(special_users, 2):
            try:
                s = SpecialNeedsForm.objects.get(user=u)
                ws7.append([
                    u.get_full_name(), str(s.birth_date), s.health_status,
                    s.education_level, s.current_city, s.current_street or '',
                    s.housing_type, s.housing_ownership, s.case_details or '',
                ])
            except SpecialNeedsForm.DoesNotExist:
                ws7.append([u.get_full_name()] + ['—'] * 8)
            _row(ws7, i)
        _auto_width(ws7)

    # ══ ورقة 8 — الكفلاء ══
    sponsor_users = [u for u in users if u.user_type == 'sponsor']
    if sponsor_users:
        ws8 = wb.create_sheet('الكفلاء')
        _hdr(ws8, ['الاسم الكامل', 'اسم المستخدم', 'البريد',
                   'الجوال', 'المهنة', 'الدولة', 'المدينة'])
        for i, u in enumerate(sponsor_users, 2):
            try:
                p = SponsorProfile.objects.get(user=u)
                ws8.append([
                    u.get_full_name(), u.username, u.email or '',
                    f'{u.phone_country or ""}{u.phone or ""}',
                    p.job or '', p.country or '', p.city or '',
                ])
            except SponsorProfile.DoesNotExist:
                ws8.append([u.get_full_name(), u.username, u.email or ''] + ['—'] * 4)
            _row(ws8, i)
        _auto_width(ws8)

    # ══ ورقة 9 — المعيلون ══
    ws9 = wb.create_sheet('المعيلون')
    _hdr(ws9, ['اسم المستفيد', 'نوعه', 'اسم المعيل', 'رقم الهوية',
               'الجنس', 'صلة القرابة', 'المهنة', 'الصحة', 'الدخل', 'المُعالون'])
    i =  2
    for u in users:
        if u.user_type == 'sponsor':
            continue
        form_id = _get_form_id(u)
        if not form_id:
            continue
        g = CurrentGuardian.objects.filter(form_type=u.user_type, form_id=form_id).first()
        if g:
            ws9.append([
                u.get_full_name(), TYPE_MAP.get(u.user_type, ''),
                g.get_full_name(), g.id_number or '',
                g.gender, g.relation, g.job,
                g.health_status, str(g.monthly_income), g.dependents,
            ])
            _row(ws9, i); i += 1
    if i == 2:
        ws9.append(['لا توجد بيانات معيلين'])
    _auto_width(ws9)

    # ══ ورقة 10 — أفراد الأسرة ══
    ws10 = wb.create_sheet('أفراد الأسرة')
    _hdr(ws10, ['اسم المستفيد', 'نوعه', 'اسم الفرد', 'رقم الهوية',
                'الجنس', 'تاريخ الميلاد', 'صلة القرابة', 'الصحة', 'التعليم'])
    i = 2
    for u in users:
        if u.user_type == 'sponsor':
            continue
        form_id = _get_form_id(u)
        if not form_id:
            continue
        members = FamilyMember.objects.filter(form_type=u.user_type, form_id=form_id)
        for m in members:
            ws10.append([
                u.get_full_name(), TYPE_MAP.get(u.user_type, ''),
                m.get_full_name(), m.id_number or '',
                m.gender, str(m.birth_date), m.relation,
                m.health_status, m.education_level or '',
            ])
            _row(ws10, i); i += 1
    if i == 2:
        ws10.append(['لا توجد أفراد أسرة'])
    _auto_width(ws10)

    # ══ إرسال الملف ══
    # إذا كان تصدير مستخدم واحد — اسمه ورقم هويته
    if len(users) == 1:
        u = users[0]
        safe_name = u.get_full_name().replace(' ', '_')
        id_num = u.id_number or u.username or str(u.pk)[:8]
        fname = f'{safe_name}_{id_num}.xlsx'
    else:
        fname = f'requests_{timezone.now().strftime("%Y%m%d_%H%M")}.xlsx'
    resp  = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    from urllib.parse import quote
    resp['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(fname)}"

    wb.save(resp)
    return resp


# ══════════════════════════════════════════════
#  طباعة
# ══════════════════════════════════════════════
@login_required
@require_GET
def print_request(request):
    user_id = request.GET.get('id', '').strip()
    if not user_id:
        return HttpResponse('يرجى تحديد معرف المستخدم', status=400)
    try:
        user   = CustomUser.objects.get(pk=user_id)
        detail = _user_detail(user)
        html   = render_to_string('admin_panel/print_request.html', {
            'user':       user,
            'detail':     detail,
            'printed_by': request.user.get_full_name() or request.user.username,
            'request':    request,
        })
        return HttpResponse(html)
    except CustomUser.DoesNotExist:
        return HttpResponse('المستخدم غير موجود', status=404)

# ══════════════════════════════════════════════
#  عداد الطلبات الجديدة (Auto-refresh)
# ══════════════════════════════════════════════
@admin_required
@require_GET
def requests_count(request):
    count = CustomUser.objects.filter(
        is_approved=False, is_active=True
    ).exclude(user_type='admin').count()
    return JsonResponse({'count': count})

"""
دوال تصدير Word و PDF لطلبات التسجيل
تُضاف إلى admin_panel/views/requests.py
"""

import io
import os
import re
import requests as http_requests
from datetime import date

from django.http import HttpResponse
from django.utils import timezone

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image as PILImage

# ══════════════════════════════════════════════
#  ثوابت الألوان
# ══════════════════════════════════════════════
HDR_COLOR = RGBColor(0x1A, 0x7A, 0x4A)   # أخضر داكن
HDR_BG    = '1A7A4A'
SUB_BG    = 'E8F5E9'
ROW_BG    = 'F1F8F4'
WHITE     = 'FFFFFF'
GRAY_BG   = 'F9F9F9'


# ══════════════════════════════════════════════
#  مساعدات Word
# ══════════════════════════════════════════════

def _set_cell_bg(cell, color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color)
    tcPr.append(shd)


def _set_rtl_para(paragraph):
    pPr  = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _cell(cell, text, bold=False, size=9, color=None, center=False, bg=None):
    """كتابة نص في خلية مع تنسيق"""
    if bg:
        _set_cell_bg(cell, bg)
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl_para(para)
    run = para.add_run(str(text) if text not in (None, '', '—', 'None') else '—')
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_row_to_table(table, data, bg=WHITE):
    """إضافة صف لجدول موجود"""
    row = table.add_row()
    for i, val in enumerate(data):
        if i < len(row.cells):
            _cell(row.cells[i], val, bg=bg)
    return row


def _make_table(doc, headers, rows_data, col_widths=None, title=None):
    """إنشاء جدول كامل"""
    if title:
        _section_title(doc, title)

    n_cols = len(headers)
    table  = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # عرض الأعمدة
    if col_widths:
        for i, w in enumerate(col_widths):
            if i < n_cols:
                for cell in table.columns[i].cells:
                    cell.width = Cm(w)

    # رأس الجدول
    hdr_row = table.rows[0]
    hdr_row.height = Cm(0.7)
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        _set_cell_bg(c, HDR_BG)
        _cell(c, h, bold=True, size=8, color=WHITE, center=True, bg=HDR_BG)

    # البيانات
    for r_idx, row_data in enumerate(rows_data):
        bg  = ROW_BG if r_idx % 2 == 0 else WHITE
        row = table.add_row()
        row.height = Cm(0.6)
        for c_idx, val in enumerate(row_data):
            if c_idx < n_cols:
                _cell(row.cells[c_idx], val, size=9, bg=bg)

    doc.add_paragraph()
    return table


def _section_title(doc, title):
    """عنوان قسم ملون"""
    p   = doc.add_paragraph()
    _set_rtl_para(p)
    run = p.add_run(f'  {title}  ')
    run.bold           = True
    run.font.size      = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name      = 'Arial'
    # خلفية خضراء للعنوان
    rPr  = run._r.get_or_add_rPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  HDR_BG)
    rPr.append(shd)
    # حدود سفلية
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:color'), HDR_BG)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def _header(doc, form_title, user_pk, user=None, detail=None):
    """رأس الصفحة مع الشعار والصورة"""
    # جدول الرأس
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # الخلية الأولى — الشعار والاسم
    c0 = table.rows[0].cells[0]
    _set_cell_bg(c0, HDR_BG)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run('رُحَمَاء')
    r0.bold = True; r0.font.size = Pt(16); r0.font.color.rgb = RGBColor.from_string(WHITE)
    r0.font.name = 'Arial'
    p0b = c0.add_paragraph('جمعية نسائم فلسطين الخيرية')
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0b.runs[0].font.size = Pt(7); p0b.runs[0].font.color.rgb = RGBColor.from_string('C8E6C9')

    # الخلية الوسطى — عنوان الاستمارة
    c1 = table.rows[0].cells[1]
    _set_cell_bg(c1, '2E7D32')
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(form_title)
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = RGBColor.from_string(WHITE)
    r1.font.name = 'Arial'

    # الخلية الثالثة — التاريخ ورقم الاستمارة
    c2 = table.rows[0].cells[2]
    _set_cell_bg(c2, HDR_BG)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f'التاريخ: {date.today().strftime("%Y/%m/%d")}')
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(WHITE); r2.font.name = 'Arial'
    reg_num = user.registration_number or str(user_pk)[:8]
    p2b = c2.add_paragraph(f'رقم الاستمارة: {reg_num}')
    p2b.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2b.runs[0].font.size = Pt(9); p2b.runs[0].font.color.rgb = RGBColor.from_string(WHITE)

    # ضبط عرض الأعمدة
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(8)
    table.columns[2].width = Cm(5)

    doc.add_paragraph()

    # صورة المستخدم إن وجدت
    if detail:
        photo_url = (detail.get('orphan') or detail.get('special') or
                     detail.get('family') or detail.get('sponsor') or {}).get('photo_url')
        if photo_url:
            _add_photo(doc, photo_url, user.get_full_name() if user else '')


def _add_photo(doc, photo_url, name=''):
    """إضافة الصورة الشخصية"""
    try:
        base_url = 'http://127.0.0.1:8000'
        full_url = base_url + photo_url if photo_url.startswith('/') else photo_url
        resp = http_requests.get(full_url, timeout=8)
        if resp.status_code == 200:
            img_bytes = io.BytesIO(resp.content)
            pil_img   = PILImage.open(img_bytes)
            img_bytes.seek(0)

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # الصورة
            c0 = table.rows[0].cells[0]
            _set_cell_bg(c0, GRAY_BG)
            p  = c0.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_bytes, width=Cm(3.5), height=Cm(4))

            # الاسم تحت الصورة
            c1 = table.rows[0].cells[1]
            _set_cell_bg(c1, SUB_BG)
            _cell(c1, f'الاسم: {name}', bold=True, size=10, center=False)

            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(14)
            doc.add_paragraph()
    except Exception:
        pass


def _addr(city, street, landmark):
    """تنسيق العنوان"""
    parts = [p for p in [city, street, landmark] if p and p != '—']
    return ' — '.join(parts) if parts else '—'


def _phone(country, number):
    """تنسيق رقم الهاتف"""
    if number and number not in ('—', '', 'None'):
        return f'{country or ""}{number}'
    return '—'


def _name_parts(full_name):
    """تقسيم الاسم الرباعي"""
    parts = (full_name or '').split()
    return (
        parts[0] if len(parts) > 0 else '—',
        parts[1] if len(parts) > 1 else '—',
        parts[2] if len(parts) > 2 else '—',
        parts[-1] if len(parts) > 3 else '—',
    )


def _footer(doc):
    """تذييل التوقيعات"""
    doc.add_paragraph()
    _section_title(doc, '✍️ التوقيعات والاعتماد')
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    labels = ['توقيع مقدم الطلب', 'توقيع المراجع', 'ختم الجمعية واعتمادها']
    for i, lbl in enumerate(labels):
        _cell(table.rows[0].cells[i], lbl, bold=True, size=9, center=True, bg=SUB_BG)
        table.rows[1].cells[i].height = Cm(1.5)
        _set_cell_bg(table.rows[1].cells[i], WHITE)

    doc.add_paragraph()
    p = doc.add_paragraph('🔒 سري وخاص — جميع البيانات محمية ومخصصة للاستخدام الرسمي فقط')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size      = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)


def _add_doc_image(doc, doc_url, doc_type):
    """إضافة وثيقة في صفحة منفصلة"""
    doc.add_page_break()

    # عنوان الوثيقة
    p   = doc.add_paragraph()
    _set_rtl_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'📎 الوثيقة: {doc_type}')
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = HDR_COLOR
    run.font.name      = 'Arial'
    doc.add_paragraph()

    ext = doc_url.split('.')[-1].lower()
    if ext == 'pdf':
        p2 = doc.add_paragraph(f'📋 ملف PDF — يرجى فتحه بشكل منفصل')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        return

    try:
        base_url = 'http://127.0.0.1:8000'
        full_url = base_url + doc_url if doc_url.startswith('/') else doc_url
        resp     = http_requests.get(full_url, timeout=10)
        if resp.status_code == 200:
            img_bytes = io.BytesIO(resp.content)
            pil_img   = PILImage.open(img_bytes)
            w, h      = pil_img.size
            ratio     = h / w if w > 0 else 1
            img_w     = Cm(15)
            img_h     = Cm(min(15 * ratio, 20))
            img_bytes.seek(0)
            p3  = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run()
            run3.add_picture(img_bytes, width=img_w, height=img_h)
    except Exception as e:
        p_err = doc.add_paragraph(f'⚠️ تعذر تحميل الصورة')
        p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_err.runs[0].font.color.rgb = RGBColor(0xE5, 0x3E, 0x3E)

def _add_payments_section(doc, user):
    """إضافة قسم المدفوعات في Word"""
    from core.models import Payment

    pays = Payment.objects.filter(
        beneficiary=user
    ).select_related('sponsor', 'created_by').order_by('-date')[:20]

    if not pays:
        return

    STATUS_MAP  = {'paid': 'مدفوعة', 'pending': 'معلّقة', 'late': 'متأخرة'}
    PAID_BY_MAP = {'sponsor': 'كافل', 'admin': 'إدارة', 'external': 'جهة خارجية'}

    rows = []
    for p in pays:
        if p.paid_by == 'sponsor' and p.sponsor:
            party = p.sponsor.get_full_name()
        elif p.paid_by == 'admin' and p.created_by:
            party = p.created_by.get_full_name() or p.created_by.username
        elif p.paid_by == 'external':
            party = p.paid_by_note or '—'
        else:
            party = '—'

        rows.append([
            str(p.date),
            f'{p.amount_ils}₪',
            f'${p.amount_usd}',
            PAID_BY_MAP.get(p.paid_by, p.paid_by),
            party,
            STATUS_MAP.get(p.status, p.status),
            p.note or '—',
        ])

    _make_table(doc,
        ['التاريخ', 'المبلغ ₪', 'المبلغ $', 'مصدر الدفع', 'الشخص/الجهة', 'الحالة', 'ملاحظة'],
        rows,
        col_widths=[2.5, 2.5, 2.5, 2.5, 3.5, 2.5, 3.0],
        title='💰 سجل المدفوعات'
    )


def _add_aids_section(doc, user):
    """إضافة قسم المساعدات في Word"""
    from core.models import Aid

    aids = Aid.objects.filter(
        beneficiary=user
    ).select_related('created_by').order_by('-date')[:20]

    if not aids:
        return

    AID_TYPES = {
        'food':      'غذائية',
        'medical':   'طبية',
        'financial': 'مالية',
        'clothing':  'ملابس',
        'furniture': 'أثاث',
        'education': 'تعليمية',
        'other':     'أخرى',
    }

    rows = []
    for a in aids:
        created_by = (a.created_by.get_full_name() or a.created_by.username) if a.created_by else 'النظام'
        rows.append([
            str(a.date),
            a.name,
            AID_TYPES.get(a.aid_type, a.aid_type),
            str(a.quantity),
            a.provider,
            created_by,
            a.note or '—',
        ])

    _make_table(doc,
        ['التاريخ', 'الاسم', 'النوع', 'الكمية', 'الجهة', 'معتمد الطلب', 'ملاحظة'],
        rows,
        col_widths=[2.5, 3.5, 2.5, 1.8, 3.0, 3.0, 3.2],
        title='🎁 سجل المساعدات'
    )


def _add_sponsor_section(doc, user):
    """إضافة قسم الكافل في Word"""
    try:
        from beneficiary.models import OrphanForm, SpecialNeedsForm, FamilyForm
        model_map = {
            'orphan':  OrphanForm,
            'special': SpecialNeedsForm,
            'family':  FamilyForm,
        }
        Model = model_map.get(user.user_type)
        if not Model:
            return
        form = Model.objects.get(user=user)
        if not (hasattr(form, 'sponsor') and form.sponsor):
            return
        sp = form.sponsor.user
        _make_table(doc,
            ['اسم الكافل', 'رقم الاستمارة', 'البريد الإلكتروني', 'رقم الجوال'],
            [[sp.get_full_name(), sp.registration_number or '—',
              sp.email, f'{sp.phone_country}{sp.phone}']],
            col_widths=[5.0, 4.5, 5.5, 4.0],
            title='🤝 بيانات الكافل'
        )
    except Exception:
        pass


# ══════════════════════════════════════════════
#  استمارة اليتيم
# ══════════════════════════════════════════════

def _build_orphan_doc(doc, user, detail):
    o = detail.get('orphan',   {}) or {}
    m = detail.get('mother',   {}) or {}
    f = detail.get('father',   {}) or {}
    g = detail.get('guardian', {}) or {}

    _header(doc, '( استمارة يتيم )', user.pk, user, detail)

    # ── البيانات الأساسية ──
    mf, mfa, mga, mla = _name_parts(m.get('full_name',''))
    gf, gfa, gga, gla = _name_parts(g.get('full_name',''))
    _make_table(doc,
        ['بيان عن', 'الاسم الأول', 'اسم الأب', 'اسم الجد', 'العائلة',
         'تاريخ الميلاد', 'رقم الهوية', 'الجنسية', 'ذكر', 'أنثى'],
        [
            ['اليتيم', user.first_name, user.father_name, user.grand_name,
             user.family_name, o.get('birth_date',''), user.id_number or '',
             user.nationality or '',
             '✓' if user.gender == 'ذكر' else '', '✓' if user.gender == 'أنثى' else ''],
            ['والدته', mf, mfa, mga, mla,
             m.get('birth_date',''), m.get('id_number',''), 'فلسطينية', '', '✓'],
            ['المعيل الحالي', gf, gfa, gga, gla,
             '', g.get('id_number',''), '', '', ''],
        ],
        col_widths=[2.2, 2.2, 2.2, 2.2, 2.2, 2.8, 2.8, 2.3, 1, 1],
        title='بيانات اليتيم الأساسية'
    )

    # ── المستوى التعليمي ──
    _make_table(doc,
        ['المستوى التعليمي', 'الصف الدراسي', 'المرحلة الدراسية', 'اسم المدرسة', 'الحالة الصحية'],
        [[o.get('education_level',''), o.get('school_grade',''),
          o.get('education_level',''), o.get('school_name',''), o.get('health_status','')]],
        col_widths=[3.5, 3, 3.5, 4, 3.5],
        title='المستوى التعليمي'
    )

    # ── بيانات الأب ──
    _make_table(doc,
        ['تاريخ الميلاد', 'تاريخ الوفاة', 'سبب الوفاة',
         'عدد الأبناء', 'إيراده قبل الوفاة', 'المعاش بعد الوفاة',
         'مستواه التعليمي', 'مهنته'],
        [[f.get('birth_date',''), f.get('death_date',''), f.get('death_reason',''),
          f.get('children_count',''), f.get('income_before','—'),
          f.get('pension_after','—'), f.get('education_level',''), f.get('job','')]],
        col_widths=[2.5, 2.5, 2.5, 2, 3, 3, 3, 2.5],
        title='بيانات والد اليتيم'
    )

    # ── بيانات الأم ──
    _make_table(doc,
        ['تاريخ الميلاد', 'تاريخ الوفاة', 'سبب الوفاة',
         'المستوى التعليمي', 'المهنة', 'الحالة الصحية', 'الدخل الشهري'],
        [[m.get('birth_date',''), m.get('death_date','—'), m.get('death_reason','—'),
          m.get('education_level',''), m.get('job',''),
          m.get('health_status',''), m.get('monthly_income','')]],
        col_widths=[2.8, 2.8, 2.8, 3.2, 3, 3, 3],
        title='بيانات والدة اليتيم'
    )

    # ── المعيل ──
    _make_table(doc,
        ['قرابته لليتيم', 'مستواه التعليمي', 'عدد من يعيلهم بما فيهم اليتيم',
         'نوع عمله', 'دخله الشهري', 'الحالة الصحية', 'المخصص التقريبي لليتيم'],
        [[g.get('relation',''), g.get('education_level',''), g.get('dependents',''),
          g.get('job',''), g.get('monthly_income',''), g.get('health_status',''), '—']],
        col_widths=[2.8, 3, 3.5, 2.8, 2.8, 2.8, 3.5],
        title='معيل اليتيم الحالي'
    )

    # ── السكن ──
    _make_table(doc,
        ['نوع السكن', 'نوع الملكية', 'مبلغ الإيجار (₪)', 'العنوان السابق', 'العنوان الحالي'],
        [[o.get('housing_type',''), o.get('housing_ownership',''),
          o.get('monthly_rent','—'),
          _addr(o.get('previous_city',''), o.get('previous_street',''), o.get('previous_landmark','')),
          _addr(o.get('current_city',''), o.get('current_street',''), o.get('current_landmark',''))]],
        col_widths=[2.5, 2.5, 2.8, 5, 5],
        title='إيضاحات السكن والعنوان'
    )

    # ── التواصل ──
    _make_table(doc,
        ['رقم جوال 1', 'رقم جوال 2', 'واتساب', 'حساب بنكي / محفظة', 'اسم صاحب الحساب'],
        [[_phone(user.phone_country, user.phone),
          _phone(getattr(user,'phone2_country',''), getattr(user,'phone2','')),
          _phone(getattr(user,'whatsapp_country',''), getattr(user,'whatsapp','')),
          '—', '—']],
        col_widths=[3.5, 3.5, 3.5, 4.5, 4.5],
        title='بيانات التواصل'
    )

    # ── أفراد الأسرة ──
    members = detail.get('members', [])
    rows = []
    for i, mem in enumerate(members, 1):
        rows.append([
            str(i), mem.get('full_name',''), mem.get('id_number',''),
            mem.get('birth_date',''), mem.get('relation',''),
            mem.get('marital_status',''), mem.get('health_status',''),
            mem.get('education_level',''), '—',
        ])
    if not rows:
        rows = [['', '', '', '', '', '', '', '', '']]
    _make_table(doc,
        ['م', 'الاسم', 'رقم الهوية', 'تاريخ الميلاد', 'صلة القرابة',
         'الحالة الاجتماعية', 'الحالة الصحية', 'التعليم', 'الدخل'],
        rows,
        col_widths=[0.8, 3.5, 2.8, 2.5, 2.2, 2.8, 2.8, 2.5, 2.2],
        title=f'بيانات الأسرة ({len(members)} فرد)'
    )

    # ── القصة ──
    if o.get('story'):
        _section_title(doc, '📖 قصة اليتيم')
        p = doc.add_paragraph(o['story'])
        _set_rtl_para(p)
        p.alignment         = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(10)

    # ── المدفوعات والمساعدات والكافل ──
    _add_payments_section(doc, user)
    _add_aids_section(doc, user)
    _add_sponsor_section(doc, user)

    _footer(doc)


# ══════════════════════════════════════════════
#  استمارة الأسرة
# ══════════════════════════════════════════════

def _build_family_doc(doc, user, detail):
    fam = detail.get('family',   {}) or {}
    w   = detail.get('wife',     {}) or {}
    g   = detail.get('guardian', {}) or {}

    _header(doc, '( استمارة أسرة )', user.pk, user, detail)

    # ── البيانات الأساسية ──
    wf, wfa, wga, wla = _name_parts(w.get('full_name',''))
    gf, gfa, gga, gla = _name_parts(g.get('full_name',''))
    _make_table(doc,
        ['بيان عن', 'الاسم الأول', 'اسم الأب', 'اسم الجد', 'العائلة',
         'تاريخ الميلاد', 'رقم الهوية', 'الجنسية', 'ذكر', 'أنثى'],
        [
            ['رب الأسرة', user.first_name, user.father_name, user.grand_name,
             user.family_name, fam.get('birth_date',''), fam.get('id_number',''),
             user.nationality or '',
             '✓' if user.gender == 'ذكر' else '', '✓' if user.gender == 'أنثى' else ''],
            ['الزوجة', wf, wfa, wga, wla,
             w.get('birth_date',''), w.get('id_number',''), '', '', '✓'],
            ['المعيل الحالي', gf, gfa, gga, gla,
             '', g.get('id_number',''), '', '', ''],
        ],
        col_widths=[2.2, 2.2, 2.2, 2.2, 2.2, 2.8, 2.8, 2.3, 1, 1],
        title='البيانات الأساسية'
    )

    # ── بيانات إضافية ──
    _make_table(doc,
        ['الحالة الاجتماعية', 'الحالة الصحية', 'المستوى التعليمي',
         'المهنة', 'عدد أفراد الأسرة', 'عدد المرضى'],
        [[fam.get('marital_status',''), fam.get('health_status',''),
          fam.get('education_level',''), fam.get('job',''),
          fam.get('members_count',''), fam.get('sick_count','')]],
        col_widths=[3, 3, 3.5, 3, 3, 3],
        title='بيانات رب الأسرة الإضافية'
    )

    # ── الوضع العام ──
    _section_title(doc, '📝 ملاحظات عن الوضع العام للأسرة')
    p = doc.add_paragraph(fam.get('general_status','') or '—')
    _set_rtl_para(p)
    p.paragraph_format.space_after = Pt(8)

    # ── السكن ──
    _make_table(doc,
        ['نوع السكن', 'نوع الملكية', 'مبلغ الإيجار (₪)', 'العنوان السابق', 'العنوان الحالي'],
        [[fam.get('housing_type',''), fam.get('housing_ownership',''),
          fam.get('monthly_rent','—'),
          _addr(fam.get('previous_city',''), fam.get('previous_street',''), fam.get('previous_landmark','')),
          _addr(fam.get('current_city',''), fam.get('current_street',''), fam.get('current_landmark',''))]],
        col_widths=[2.5, 2.5, 2.8, 5, 5],
        title='إيضاحات السكن والعنوان'
    )

    # ── التواصل ──
    _make_table(doc,
        ['رقم جوال 1', 'رقم جوال 2', 'واتساب', 'حساب بنكي / محفظة', 'اسم صاحب الحساب'],
        [[_phone(user.phone_country, user.phone),
          _phone(getattr(user,'phone2_country',''), getattr(user,'phone2','')),
          _phone(getattr(user,'whatsapp_country',''), getattr(user,'whatsapp','')),
          '—', '—']],
        col_widths=[3.5, 3.5, 3.5, 4.5, 4.5],
        title='بيانات التواصل'
    )

    # ── المعيل ──
    _make_table(doc,
        ['قرابته', 'مستواه التعليمي', 'عدد من يعيلهم',
         'نوع عمله', 'دخله الشهري', 'الحالة الصحية'],
        [[g.get('relation',''), g.get('education_level',''), g.get('dependents',''),
          g.get('job',''), g.get('monthly_income',''), g.get('health_status','')]],
        col_widths=[2.8, 3.2, 3, 3, 3, 3],
        title='بيانات المعيل الحالي'
    )

    # ── أفراد الأسرة ──
    members = detail.get('members', [])
    rows = []
    for i, mem in enumerate(members, 1):
        rows.append([
            str(i), mem.get('full_name',''), mem.get('id_number',''),
            mem.get('birth_date',''), mem.get('relation',''),
            mem.get('marital_status',''), mem.get('health_status',''),
            mem.get('education_level',''), '—',
        ])
    if not rows:
        rows = [['', '', '', '', '', '', '', '', '']]
    _make_table(doc,
        ['م', 'الاسم', 'رقم الهوية', 'تاريخ الميلاد', 'صلة القرابة',
         'الحالة الاجتماعية', 'الحالة الصحية', 'التعليم', 'الدخل'],
        rows,
        col_widths=[0.8, 3.5, 2.8, 2.5, 2.2, 2.8, 2.8, 2.5, 2.2],
        title=f'بيانات الأسرة ({len(members)} فرد)'
    )

    # ── المدفوعات والمساعدات والكافل ──
    _add_payments_section(doc, user)
    _add_aids_section(doc, user)
    _add_sponsor_section(doc, user)

    _footer(doc)


# ══════════════════════════════════════════════
#  استمارة ذوو الاحتياجات
# ══════════════════════════════════════════════

def _build_special_doc(doc, user, detail):
    s = detail.get('special',  {}) or {}
    g = detail.get('guardian', {}) or {}

    _header(doc, '( استمارة ذوي الاحتياجات الخاصة )', user.pk, user, detail)

    # ── البيانات الأساسية ──
    gf, gfa, gga, gla = _name_parts(g.get('full_name',''))
    _make_table(doc,
        ['بيان عن', 'الاسم الأول', 'اسم الأب', 'اسم الجد', 'العائلة',
         'تاريخ الميلاد', 'رقم الهوية', 'الجنسية', 'ذكر', 'أنثى'],
        [
            ['المريض', user.first_name, user.father_name, user.grand_name,
             user.family_name, s.get('birth_date',''), user.id_number or '',
             user.nationality or '',
             '✓' if user.gender == 'ذكر' else '', '✓' if user.gender == 'أنثى' else ''],
            ['المعيل الحالي', gf, gfa, gga, gla,
             '', g.get('id_number',''), '', '', ''],
        ],
        col_widths=[2.2, 2.2, 2.2, 2.2, 2.2, 2.8, 2.8, 2.3, 1, 1],
        title='البيانات الأساسية'
    )

    # ── التعليم والصحة ──
    _make_table(doc,
        ['المستوى التعليمي', 'الصف الدراسي', 'اسم المدرسة', 'الحالة الصحية'],
        [[s.get('education_level',''), s.get('school_grade',''),
          s.get('school_name',''), s.get('health_status','')]],
        col_widths=[4, 3, 5, 4.5],
        title='المستوى التعليمي والصحي'
    )

    # ── تفاصيل الحالة ──
    _section_title(doc, '🏥 تفاصيل الحالة الصحية')
    p = doc.add_paragraph(s.get('case_details','') or '—')
    _set_rtl_para(p)
    p.paragraph_format.space_after = Pt(8)

    # ── المعيل ──
    _make_table(doc,
        ['قرابته للمريض', 'مستواه التعليمي', 'عدد من يعيلهم بما فيهم المريض',
         'نوع عمله', 'دخله الشهري', 'الحالة الصحية', 'المخصص التقريبي للمريض'],
        [[g.get('relation',''), g.get('education_level',''), g.get('dependents',''),
          g.get('job',''), g.get('monthly_income',''), g.get('health_status',''), '—']],
        col_widths=[2.8, 3, 3.5, 2.8, 2.8, 2.8, 3.5],
        title='بيانات المعيل الحالي'
    )

    # ── السكن ──
    _make_table(doc,
        ['نوع السكن', 'نوع الملكية', 'مبلغ الإيجار (₪)', 'العنوان السابق', 'العنوان الحالي'],
        [[s.get('housing_type',''), s.get('housing_ownership',''),
          s.get('monthly_rent','—'),
          _addr(s.get('previous_city',''), s.get('previous_street',''), s.get('previous_landmark','')),
          _addr(s.get('current_city',''), s.get('current_street',''), s.get('current_landmark',''))]],
        col_widths=[2.5, 2.5, 2.8, 5, 5],
        title='إيضاحات السكن والعنوان'
    )

    # ── التواصل ──
    _make_table(doc,
        ['رقم جوال 1', 'رقم جوال 2', 'واتساب', 'حساب بنكي / محفظة', 'اسم صاحب الحساب'],
        [[_phone(user.phone_country, user.phone),
          _phone(getattr(user,'phone2_country',''), getattr(user,'phone2','')),
          _phone(getattr(user,'whatsapp_country',''), getattr(user,'whatsapp','')),
          '—', '—']],
        col_widths=[3.5, 3.5, 3.5, 4.5, 4.5],
        title='بيانات التواصل'
    )

    # ── أفراد الأسرة ──
    members = detail.get('members', [])
    rows = []
    for i, mem in enumerate(members, 1):
        rows.append([
            str(i), mem.get('full_name',''), mem.get('id_number',''),
            mem.get('birth_date',''), mem.get('relation',''),
            mem.get('marital_status',''), mem.get('health_status',''), '—',
        ])
    if not rows:
        rows = [['', '', '', '', '', '', '', '']]
    _make_table(doc,
        ['م', 'الاسم', 'رقم الهوية', 'تاريخ الميلاد', 'صلة القرابة',
         'الحالة الاجتماعية', 'الحالة الصحية', 'الدخل'],
        rows,
        col_widths=[0.8, 3.5, 2.8, 2.5, 2.5, 3, 3, 2.5],
        title=f'بيانات الأسرة ({len(members)} فرد)'
    )

    # ── المدفوعات والمساعدات والكافل ──
    _add_payments_section(doc, user)
    _add_aids_section(doc, user)
    _add_sponsor_section(doc, user)

    _footer(doc)


# ══════════════════════════════════════════════
#  استمارة الكافل
# ══════════════════════════════════════════════

def _build_sponsor_doc(doc, user, detail):
    sp = detail.get('sponsor', {}) or {}

    _header(doc, '( استمارة كافل )', user.pk, user, detail)

    _make_table(doc,
        ['الاسم الأول', 'اسم الأب', 'اسم الجد', 'اسم العائلة', 'الجنسية', 'الجنس'],
        [[user.first_name, user.father_name, user.grand_name,
          user.family_name, user.nationality or '', user.gender or '']],
        col_widths=[3, 3, 3, 3, 3, 2.5],
        title='البيانات الشخصية'
    )

    _make_table(doc,
        ['اسم المستخدم', 'البريد الإلكتروني', 'تاريخ التسجيل'],
        [[user.username, user.email, user.date_joined.strftime('%Y/%m/%d')]],
        col_widths=[5, 8, 4.5],
        title='بيانات الحساب'
    )

    _make_table(doc,
        ['المهنة', 'الدولة', 'المدينة'],
        [[sp.get('job','—'), sp.get('country','—'), sp.get('city','—')]],
        col_widths=[6, 6, 5.5],
        title='البيانات المهنية والجغرافية'
    )

    _make_table(doc,
        ['رقم جوال 1', 'رقم جوال 2', 'واتساب'],
        [[_phone(user.phone_country, user.phone),
          _phone(getattr(user,'phone2_country',''), getattr(user,'phone2','')),
          _phone(getattr(user,'whatsapp_country',''), getattr(user,'whatsapp',''))]],
        col_widths=[6, 6, 5.5],
        title='بيانات التواصل'
    )

    _section_title(doc, '📝 ملاحظات المراجع')
    for _ in range(5):
        p = doc.add_paragraph()
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bot)
        pPr.append(pBdr)

    _footer(doc)


# ══════════════════════════════════════════════
#  دالة التصدير — Word
# ══════════════════════════════════════════════

from .decorators import admin_required
from django.views.decorators.http import require_GET
from core.models import CustomUser


@admin_required
@require_GET
def export_word(request):
    from urllib.parse import quote

    user_id = request.GET.get('id', '').strip()
    if not user_id:
        return HttpResponse('يرجى تحديد المستخدم', status=400)

    try:
        user   = CustomUser.objects.get(pk=user_id)
        detail = _user_detail(user)
    except CustomUser.DoesNotExist:
        return HttpResponse('المستخدم غير موجود', status=404)

    doc = Document()

    section               = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    style           = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    ut = user.user_type
    if ut == 'orphan':
        _build_orphan_doc(doc, user, detail)
    elif ut == 'family':
        _build_family_doc(doc, user, detail)
    elif ut == 'special':
        _build_special_doc(doc, user, detail)
    elif ut == 'sponsor':
        _build_sponsor_doc(doc, user, detail)

    docs = detail.get('docs', [])
    if docs:
        doc.add_page_break()
        p = doc.add_paragraph('الوثائق المرفقة')
        p.alignment              = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold           = True
        p.runs[0].font.size      = Pt(14)
        p.runs[0].font.color.rgb = HDR_COLOR
        p.runs[0].font.name      = 'Arial'
        for d in docs:
            _add_doc_image(doc, d['url'], d['type'])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    name   = user.get_full_name().replace(' ', '_')
    id_num = user.id_number or user.username or str(user.pk)[:8]
    fname  = f'{name}_{id_num}.docx'

    resp = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    resp['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(fname)}"
    return resp


# ══════════════════════════════════════════════
#  دالة التصدير — PDF
# ══════════════════════════════════════════════

@login_required
@require_GET
def export_pdf(request):
    from urllib.parse import quote
    from django.template.loader import render_to_string
    import weasyprint
    import logging
    logging.getLogger('weasyprint').setLevel(logging.ERROR)
    logging.getLogger('fontTools').setLevel(logging.ERROR)

    user_id = request.GET.get('id', '').strip()
    if not user_id:
        return HttpResponse('يرجى تحديد المستخدم', status=400)

    try:
        user   = CustomUser.objects.get(pk=user_id)
        detail = _user_detail(user)
    except CustomUser.DoesNotExist:
        return HttpResponse('المستخدم غير موجود', status=404)

    html_content = render_to_string('admin_panel/pdf_request.html', {
        'user':       user,
        'detail':     detail,
        'printed_by': request.user.get_full_name() or request.user.username,
        'request':    request,
    })

    pdf_file = weasyprint.HTML(
        string=html_content,
        base_url=request.build_absolute_uri('/')
    ).write_pdf()

    name   = user.get_full_name().replace(' ', '_')
    id_num = user.id_number or user.username or str(user.pk)[:8]
    fname  = f'{name}_{id_num}.pdf'

    resp = HttpResponse(pdf_file, content_type='application/pdf')
    resp['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(fname)}"
    return resp


@require_GET
def check_field(request):
    field      = request.GET.get('field', '').strip()
    value      = request.GET.get('value', '').strip()
    exclude_id = request.GET.get('exclude_id', '').strip()

    if not field or not value:
        return JsonResponse({'available': True})

    # حقول موجودة فقط في النماذج وليس في CustomUser
    FORM_ONLY_FIELDS = ['phone2', 'phone2_country', 'whatsapp_country']
    if field in FORM_ONLY_FIELDS:
        return JsonResponse({'available': True})

    try:
        qs = CustomUser.objects.filter(**{field: value})
        if exclude_id:
            qs = qs.exclude(pk=exclude_id)
        return JsonResponse({'available': not qs.exists()})
    except Exception:
        return JsonResponse({'available': True})